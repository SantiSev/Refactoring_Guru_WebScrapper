import os
import time

from PIL import Image
from docx import Document
from docx2pdf import convert

from scraper import Decorator
from scraper.endpoints import design_patterns
from scraper.webScrapper import WebScraper

DOCX_NAME = 'refactoring_guru_summary.docx'
base_url = 'https://refactoring.guru/es/design-patterns/'


def create_cover(document):

    # add a picture to the document
    picture_path = 'scraper/cover.png'
    picture = document.add_picture(picture_path)

    # position the picture at the top left
    picture.left = 100
    picture.top = 100
    document.add_page_break()


def create_summary():

    document = Document()

    for i in range(3):
        add_page(document,design_patterns[i])
        print(design_patterns[i])
        time.sleep(2)


    # save the document as a .docx file
    document.save(DOCX_NAME)
    # convert the .docx file to a PDF
    convert(DOCX_NAME)



def add_page(document, pattern):

    url = base_url + pattern
    scraper = WebScraper(url)

    document.add_heading(f'{pattern}', 0)
    # section intent
    Decorator.add_title(document, "Proposito")
    Decorator.add_parragraph(document, scraper.get_text("section intent"),3)
    Decorator.add_image(document,scraper.get_image('section intent'), 0.3)
    # section problem
    Decorator.add_title(document, "Problema")
    Decorator.add_parragraph(document, scraper.get_text("section problem"))
    Decorator.add_image(document, scraper.get_image('section problem'), 0.3)
    # section solution
    Decorator.add_title(document, "Solucion")
    Decorator.add_parragraph(document, scraper.get_text("section solution"))
    Decorator.add_image(document, scraper.get_image('section solution'), 0.3)
    # # sector estructura
    # Decorator.add_title(document, "Solucion")
    # Decorator.add_parragraph(document, scraper.get_text("section solution"))
    # Decorator.add_image(document, scraper.get_image('section solution'), 0.5)
    # # sector pseudocodigo
    # Decorator.add_title(document, "Solucion")
    # Decorator.add_parragraph(document, scraper.get_text("section solution"))
    # Decorator.add_image(document, scraper.get_image('section solution'), 0.5)




    document.add_page_break()

