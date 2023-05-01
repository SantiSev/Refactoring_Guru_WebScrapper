import os

import docx
from PIL import Image, ImageOps
from docx.shared import Pt
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4

PAGE_WIDTH, PAGE_HEIGHT = A4  # 595.2755905511812 , 841.8897637795277
styles = getSampleStyleSheet()


def create_cover(pdf):
    pdf.drawImage('refactoringGuruScrapper\cover.png', 0, 0, PAGE_WIDTH, PAGE_HEIGHT, mask='auto')
    pdf.showPage()


def add_parragraph(document, textList,size=9 ):
    paragraph_text = "\n".join(p.text for p in textList)
    paragraph = document.add_paragraph(paragraph_text)
    paragraph.style.font.size = Pt(size)
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = docx.shared.Pt(2)
    # paragraph_format.space_after = docx.shared.Pt(0)


def add_title(document, title):
    subtitle = document.add_paragraph(title)
    subtitle.style = 'Subtitle'
    subtitle_font = subtitle.add_run().font
    subtitle_font.name = 'Arial'
    subtitle_font.size = docx.shared.Pt(16)
    subtitle.add_run(':')
    for run in subtitle.runs:
        run.underline = True
        run.bold = True
        run.font.color.rgb = docx.shared.RGBColor(0x80, 0x80, 0x80)
    subtitle.paragraph_format.space_before = docx.shared.Pt(6)
    subtitle.paragraph_format.space_after = docx.shared.Pt(0)


def add_image(document, image_path, percentage):
    scale_image(image_path, percentage)
    add_border(image_path)
    document.add_picture(image_path)
    os.remove(image_path)


def add_border(image_path):
    # Open the image file
    image = Image.open(image_path)

    # Convert the image to RGB mode
    if image.mode == 'RGBA':
        image = image.convert('RGB')

    # Add a black border with a width of 10 pixels
    bordered_image = ImageOps.expand(image, border=10, fill='black')

    # Save the bordered image
    bordered_image.save('bordered_image.jpg')

def scale_image(image_path, percentage):
    # Open the image file
    with Image.open(image_path) as image:
        # Get the current size of the image
        width, height = image.size

        # Scale down the image to the desired percentage of its original size
        new_width = int(width * percentage)
        new_height = int(height * percentage)
        resized_image = image.resize((new_width, new_height), resample=Image.LANCZOS)

        # Save the scaled down image
        resized_image.save(image_path, quality=95)